import docx
from docx.shared import Pt
from docx.shared import Inches
import io

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

doc = docx.Document()

paragraph = doc.add_paragraph()

run = paragraph.add_run(getText('fontsss.docx'))
font = run.font
font.name = 'bwhebb'
doc.save('HelloWorld.docx')